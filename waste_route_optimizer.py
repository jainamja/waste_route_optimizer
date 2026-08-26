import os
import re
import pandas as pd
from flask import Flask, request, render_template, redirect, url_for, jsonify, send_file, flash
from docx import Document
from aco_vrp import ACO_VRP
import requests
from functools import lru_cache
import time
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from flask_sqlalchemy import SQLAlchemy

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.secret_key = 'supersecretkey'

# Automatically use PostgreSQL if DATABASE_URL is set, otherwise fall back to local SQLite
database_url = os.environ.get('DATABASE_URL')
if database_url and database_url.startswith('postgres://'):
    database_url = database_url.replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_DATABASE_URI'] = database_url or 'sqlite:///routes.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Models
class User(UserMixin, db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(20), default='user')

class Customer(db.Model):
    __tablename__ = 'customers'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255))
    phone = db.Column(db.String(50))
    address = db.Column(db.Text)
    location_url = db.Column(db.Text)
    lat = db.Column(db.Float)
    lng = db.Column(db.Float)
    status = db.Column(db.String(50), default='PENDING')
    truck_id = db.Column(db.Integer)
    stop_number = db.Column(db.Integer)

class Metadata(db.Model):
    __tablename__ = 'metadata_store'
    key = db.Column(db.String(100), primary_key=True)
    value = db.Column(db.Text)

with app.app_context():
    db.create_all()
    # Auto-create a default admin user
    if not User.query.filter_by(username='admin').first():
        admin = User(username='admin', password_hash=generate_password_hash('password123'), role='admin')
        db.session.add(admin)
        db.session.commit()

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        if Customer.query.first(): return redirect(url_for('dashboard'))
        else: return redirect(url_for('index'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            if Customer.query.first(): return redirect(url_for('dashboard'))
            else: return redirect(url_for('index'))
        else:
            flash('Invalid username or password')
            
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        if Customer.query.first(): return redirect(url_for('dashboard'))
        else: return redirect(url_for('index'))
        
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if User.query.filter_by(username=username).first():
            flash('Username already exists')
        else:
            hashed_pw = generate_password_hash(password)
            new_user = User(username=username, password_hash=hashed_pw)
            db.session.add(new_user)
            db.session.commit()
            flash('Registration successful! Please log in.', 'success')
            return redirect(url_for('login'))
            
    return render_template('register.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@lru_cache(maxsize=100)
def resolve_gmaps_url(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        res = requests.get(url, allow_redirects=True, headers=headers, timeout=10)
        
        pin_match = re.search(r'!3d(-?\d+\.\d+)!4d(-?\d+\.\d+)', res.url)
        if pin_match: return float(pin_match.group(1)), float(pin_match.group(2))
        q_match = re.search(r'q=(-?\d+\.\d+),(-?\d+\.\d+)', res.url)
        if q_match: return float(q_match.group(1)), float(q_match.group(2))
        vp_match = re.search(r'@(-?\d+\.\d+),(-?\d+\.\d+)', res.url)
        if vp_match: return float(vp_match.group(1)), float(vp_match.group(2))
            
        meta_match = re.search(r'center=(-?\d+\.\d+)%2C(-?\d+\.\d+)', res.text)
        if meta_match: return float(meta_match.group(1)), float(meta_match.group(2))
            
        js_match = re.search(r'\[(2[2-4]\.\d+),([7][1-4]\.\d+)\]', res.text)
        if js_match: return float(js_match.group(1)), float(js_match.group(2))
    except Exception as e:
        print(f"Error resolving {url}: {e}")
    return None, None

def extract_lat_lng(coord_str):
    if pd.isna(coord_str): return None, None
    match = re.search(r'(-?\d+(?:\.\d+)?)\s*,\s*(-?\d+(?:\.\d+)?)', str(coord_str).strip())
    return (float(match.group(1)), float(match.group(2))) if match else (None, None)

def read_data_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    df = None
    if ext == '.xlsx':
        df = pd.read_excel(filepath)
    elif ext == '.docx':
        doc = Document(filepath)
        if doc.tables:
            table = doc.tables[0]
            data = [[cell.text for cell in row.cells] for row in table.rows]
            df = pd.DataFrame(data[1:], columns=data[0])
    
    if df is None or df.empty: return []

    cols = [str(c).lower() for c in df.columns]
    df.columns = cols
    
    coord_col = next((c for c in cols if 'coord' in c or ('lat' in c and 'lon' in c) or ('lat' in c and 'lng' in c)), None)
    lat_col = next((c for c in cols if 'lat' in c), None) if not coord_col else None
    lng_col = next((c for c in cols if 'lng' in c or 'lon' in c), None) if not coord_col else None

    name_col = next((c for c in cols if 'name' in c), cols[0])
    phone_col = next((c for c in cols if 'phone' in c), cols[1] if len(cols)>1 else None)
    address_col = next((c for c in cols if 'address' in c), cols[2] if len(cols)>2 else None)
    loc_url_col = next((c for c in cols if 'location' in c and c != coord_col), None)

    urls_to_resolve = set()
    for _, row in df.iterrows():
        loc_url_raw = row.get(loc_url_col) if loc_url_col and not pd.isna(row.get(loc_url_col)) else ""
        if loc_url_raw:
            url_match = re.search(r'(https?://[^\s]+)', str(loc_url_raw))
            if url_match: urls_to_resolve.add(url_match.group(1))
                
    resolved_urls = {}
    if urls_to_resolve:
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            future_to_url = {executor.submit(resolve_gmaps_url, url): url for url in urls_to_resolve}
            for future in concurrent.futures.as_completed(future_to_url):
                resolved_urls[future_to_url[future]] = future.result() if future.result() else (None, None)

    customers = []
    for idx, row in df.iterrows():
        lat, lng = None, None
        loc_url_clean = ""
        loc_url_raw = row.get(loc_url_col) if loc_url_col and not pd.isna(row.get(loc_url_col)) else ""
        if loc_url_raw:
            url_match = re.search(r'(https?://[^\s]+)', str(loc_url_raw))
            if url_match:
                loc_url_clean = url_match.group(1)
                if loc_url_clean in resolved_urls and resolved_urls[loc_url_clean][0]:
                    lat, lng = resolved_urls[loc_url_clean]
                
        if lat is None or lng is None:
            if coord_col: lat, lng = extract_lat_lng(row[coord_col])
            elif lat_col and lng_col: lat, lng = row[lat_col], row[lng_col]
                
        if lat is None or lng is None:
            address = row.get(address_col) if address_col and not pd.isna(row.get(address_col)) else ""
            if address:
                try:
                    from geopy.geocoders import ArcGIS
                    location = ArcGIS().geocode(address)
                    if location: lat, lng = location.latitude, location.longitude
                except: pass
            
        try:
            lat, lng = float(lat), float(lng)
            if not pd.isna(lat) and not pd.isna(lng):
                customers.append({
                    'id': idx + 1,
                    'name': row.get(name_col) if name_col and not pd.isna(row.get(name_col)) else f"Customer {idx+1}",
                    'phone': row.get(phone_col) if phone_col and not pd.isna(row.get(phone_col)) else "",
                    'address': row.get(address_col) if address_col and not pd.isna(row.get(address_col)) else "",
                    'location_url': loc_url_clean,
                    'lat': lat,
                    'lng': lng,
                    'status': 'PENDING'
                })
        except (ValueError, TypeError):
            pass
            
    return customers

@app.route('/')
@login_required
def index():
    return render_template('select_start_point.html')

@app.route('/upload', methods=['POST'])
@login_required
def upload():
    if 'file' not in request.files: return "No file part", 400
    file = request.files['file']
    if file.filename == '': return "No selected file", 400
        
    start_lat = request.form.get('start_lat')
    start_lng = request.form.get('start_lng')
    end_lat = request.form.get('end_lat')
    end_lng = request.form.get('end_lng')
    num_trucks = int(request.form.get('num_trucks', 3))
    
    if not all([start_lat, start_lng, end_lat, end_lng]): return "Missing coordinates", 400
        
    start_coord = (float(start_lat), float(start_lng))
    end_coord = (float(end_lat), float(end_lng))
    
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filepath)
    
    customers_data = read_data_file(filepath)
    if not customers_data: return "No valid customer data found", 400
        
    aco = ACO_VRP(start_coord, end_coord, customers_data, num_trucks=num_trucks)
    routes, _ = aco.run()
    
    for truck_idx, route in enumerate(routes):
        for stop_num, customer_id in enumerate(route):
            for c in customers_data:
                if c['id'] == customer_id:
                    c['truck'] = truck_idx + 1
                    c['stop_number'] = stop_num + 1

    # Clear old data
    Customer.query.delete()
    Metadata.query.delete()
    
    # Insert new data
    m1 = Metadata(key='start_coord', value=f"{start_coord[0]},{start_coord[1]}")
    m2 = Metadata(key='end_coord', value=f"{end_coord[0]},{end_coord[1]}")
    db.session.add_all([m1, m2])
    
    for c in customers_data:
        new_cust = Customer(
            id=c['id'], name=c['name'], phone=c['phone'], address=c['address'],
            location_url=c['location_url'], lat=c['lat'], lng=c['lng'],
            status=c.get('status', 'PENDING'), truck_id=c.get('truck'), stop_number=c.get('stop_number')
        )
        db.session.add(new_cust)
        
    db.session.commit()

    return redirect(url_for('dashboard'))

@app.route('/dashboard')
@login_required
def dashboard():
    if not Customer.query.first():
        return redirect(url_for('index'))
    return render_template('live_dashboard.html')

@app.route('/api/data')
@login_required
def get_data():
    metadata_rows = Metadata.query.all()
    metadata = {row.key: row.value for row in metadata_rows}
    
    start_coord = [float(x) for x in metadata['start_coord'].split(',')] if 'start_coord' in metadata else None
    end_coord = [float(x) for x in metadata['end_coord'].split(',')] if 'end_coord' in metadata else None
        
    customers_db = Customer.query.order_by(Customer.truck_id.asc(), Customer.stop_number.asc()).all()
    
    customers = []
    routes_dict = {}
    
    for r in customers_db:
        customers.append({
            'id': r.id, 'name': r.name, 'phone': r.phone, 'address': r.address,
            'location_url': r.location_url, 'lat': r.lat, 'lng': r.lng,
            'status': r.status, 'truck': r.truck_id, 'stop_number': r.stop_number
        })
        t_id = r.truck_id
        if t_id not in routes_dict: routes_dict[t_id] = []
        routes_dict[t_id].append(r.id)
        
    routes = [routes_dict[t_id] for t_id in sorted(routes_dict.keys())]

    return jsonify({
        'customers': customers,
        'routes': routes,
        'start_coord': start_coord,
        'end_coord': end_coord
    })

@app.route('/api/mark_completed/<int:customer_id>', methods=['POST'])
@login_required
def mark_completed(customer_id):
    customer = db.session.get(Customer, customer_id)
    if customer:
        customer.status = 'COMPLETED'
        db.session.commit()
    return jsonify({'success': True, 'status': 'COMPLETED'})

@app.route('/download_excel')
@login_required
def download_excel():
    customers = Customer.query.order_by(Customer.truck_id.asc(), Customer.stop_number.asc()).all()
    if not customers: return "No data", 400
    
    data = [{
        'truck': c.truck_id, 'stop_number': c.stop_number, 'name': c.name,
        'phone': c.phone, 'address': c.address, 'location_url': c.location_url,
        'lat': c.lat, 'lng': c.lng, 'status': c.status
    } for c in customers]
    
    df = pd.DataFrame(data)
    export_path = os.path.join(app.config['UPLOAD_FOLDER'], 'optimized_routes.xlsx')
    df.to_excel(export_path, index=False)
    
    return send_file(export_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, port=5000)
